#!/usr/bin/env python3
"""
Build the submittable copies of the CMST 387 coursework documents.

Markdown under docs/ is the single source of truth. This script regenerates
everything derived from it:

    python3 docs/tools/build_docs.py            # rebuild both .docx files
    python3 docs/tools/build_docs.py --html     # also emit .html for Google Drive import
    python3 docs/tools/build_docs.py --figures  # regenerate figure PNGs from the SVGs

Dependencies: python-docx (already present on this machine). No Node, no
node_modules, no network access required.

Why headless Chrome for SVG -> PNG: cairosvg fails here (no system libcairo),
rsvg-convert is not installed, and macOS `qlmanage -t` center-crops and clips
wide diagrams. Chrome renders the SVG at its exact viewBox size at 2x scale.
"""

import argparse
import html as html_mod
import os
import re
import shutil
import subprocess
import sys
import xml.etree.ElementTree as ET
from pathlib import Path

DOCS = Path(__file__).resolve().parent.parent
FIGURES = DOCS / "figures"

CHROME_CANDIDATES = [
    "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome",
    "/Applications/Chromium.app/Contents/MacOS/Chromium",
    "/usr/bin/google-chrome",
    "/usr/bin/chromium",
]

COURSE = "CMST 387 6300 — Principles of Web Design and Technology III"
STUDENT = "Gilbert Anderson"

# Figures inserted into the proposal, in order, with their captions.
PROPOSAL_FIGURES = [
    ("wireframe-home.png", "Figure 1. Home page wireframe (desktop and mobile)."),
    ("wireframe-interior.png", "Figure 2. Interior topic page wireframe (desktop and mobile)."),
    ("wireframe-contact.png", "Figure 3. Contact page wireframe with responsive form (desktop and mobile)."),
    ("sitemap-storyboard.png", "Figure 4. Site map and user flow diagram."),
]

DOCUMENTS = [
    {
        "md": "project1-proposal.md",
        "docx": "Project1-Proposal-Anderson.docx",
        "html": "project1-proposal.html",
        "title": "Project 1: Web Project Proposal",
        "figures": True,
    },
    {
        "md": "project1-reflection.md",
        "docx": "Project1-Reflection-Anderson.docx",
        "html": "project1-reflection.html",
        "title": "Project 1: Project Reflection",
        "figures": False,
    },
    {
        "md": "course-log.md",
        "docx": "CMST387-Course-Log-Anderson.docx",
        "html": "course-log.html",
        "title": "CMST 387 — Running Course Log",
        "figures": False,
        "optional": True,      # only built once the log exists
    },
]


# --------------------------------------------------------------------------
# Shared markdown parsing
# --------------------------------------------------------------------------

def parse_blocks(md_text):
    """Turn markdown into a flat list of (kind, payload) blocks.

    kinds: heading (level, text) | para (text) | bullet (text) |
           table (rows) | rule | fence (text)
    """
    # Strip HTML comments first: an entry template inside <!-- --> would
    # otherwise be parsed as real headings and land in the output.
    md_text = re.sub(r"<!--.*?-->", "", md_text, flags=re.S)

    lines = md_text.split("\n")
    blocks = []
    i = 0
    in_refs = False

    while i < len(lines):
        raw = lines[i]
        s = raw.strip()

        # fenced code
        if s.startswith("```"):
            body = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                body.append(lines[i])
                i += 1
            blocks.append(("fence", "\n".join(body)))
            i += 1
            continue

        # table
        if s.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{2,}:?", c) for c in cells):
                    rows.append(cells)
                i += 1
            blocks.append(("table", rows))
            continue

        if s.startswith("#"):
            level = len(s) - len(s.lstrip("#"))
            text = s.lstrip("#").strip()
            in_refs = text.lower() == "references"
            blocks.append(("heading", (level, text)))
            i += 1
            continue

        if s == "---":
            blocks.append(("rule", None))
            i += 1
            continue

        if s.startswith("- "):
            body = s[2:]
            j = i + 1
            # absorb hanging continuation lines
            while (j < len(lines) and lines[j].startswith("  ")
                   and lines[j].strip() and not lines[j].strip().startswith("- ")):
                body += " " + lines[j].strip()
                j += 1
            blocks.append(("bullet", body))
            i = j
            continue

        if not s:
            i += 1
            continue

        # paragraph: join wrapped lines, but start a new paragraph at a bold label
        para = [s]
        j = i + 1
        while (j < len(lines) and lines[j].strip()
               and not lines[j].strip().startswith(("#", "-", "|", "**", "```"))
               and lines[j].strip() != "---"):
            para.append(lines[j].strip())
            j += 1
        blocks.append(("refpara" if in_refs else "para", " ".join(para)))
        i = j

    return blocks


def split_inline(text):
    """Split markdown inline formatting into (text, bold, italic, code) runs.

    Bold is matched non-greedily so **bold wrapping *italic*** works; italic is
    matched only outside bold markers so it does not eat the bold delimiters.
    """
    tokens = []
    pattern = re.compile(
        r"(\[[^\]]+\]\([^)]+\)|\*\*.+?\*\*|`[^`]+`|(?<![\*\w])\*[^*]+?\*(?![\*\w]))")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            tokens.append((text[pos:m.start()], False, False, False, None))
        tok = m.group(0)
        if tok.startswith("["):
            label, href = re.match(r"\[([^\]]+)\]\(([^)]+)\)", tok).groups()
            tokens.append((label, False, False, False, href))
        elif tok.startswith("**"):
            inner = tok[2:-2]
            # allow *italic* nested inside bold
            sub = re.split(r"(\*[^*]+\*)", inner)
            for part in sub:
                if not part:
                    continue
                if part.startswith("*") and part.endswith("*"):
                    tokens.append((part[1:-1], True, True, False, None))
                else:
                    tokens.append((part, True, False, False, None))
        elif tok.startswith("`"):
            tokens.append((tok[1:-1], False, False, True, None))
        else:
            tokens.append((tok[1:-1], False, True, False, None))
        pos = m.end()
    if pos < len(text):
        tokens.append((text[pos:], False, False, False, None))
    return tokens


# --------------------------------------------------------------------------
# DOCX output
# --------------------------------------------------------------------------

def build_docx(blocks, out_path, title, figures=False):
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11.5)
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing = 1.15

    def add_runs(par, text):
        for chunk, bold, italic, code, href in split_inline(text):
            run = par.add_run(chunk)
            if href:
                run.underline = True
            run.bold = bold
            run.italic = italic
            if code:
                run.font.name = "Consolas"
                run.font.size = Pt(10.5)

    # ---- title block ----
    h = doc.add_paragraph()
    r = h.add_run(title)
    r.bold = True
    r.font.size = Pt(20)
    for label, value in (("Course: ", COURSE), ("Student: ", STUDENT)):
        p = doc.add_paragraph()
        p.add_run(label).bold = True
        p.add_run(value)
    doc.add_paragraph()

    seen_title = False
    for kind, payload in blocks:
        if kind == "heading":
            level, text = payload
            if level == 1 and not seen_title:
                seen_title = True          # already rendered as the title block
                continue
            doc.add_heading(text, level=min(level, 4))

        elif kind in ("para", "refpara"):
            p = doc.add_paragraph()
            add_runs(p, payload)
            if kind == "refpara":          # APA hanging indent
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.first_line_indent = Inches(-0.5)

        elif kind == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, payload)

        elif kind == "fence":
            p = doc.add_paragraph()
            run = p.add_run(payload)
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)

        elif kind == "table":
            rows = payload
            if not rows:
                continue
            t = doc.add_table(rows=len(rows), cols=len(rows[0]))
            t.style = "Table Grid"
            for ri, row in enumerate(rows):
                for ci, cell in enumerate(row):
                    cell_par = t.cell(ri, ci).paragraphs[0]
                    add_runs(cell_par, cell)
                    if ri == 0:
                        for run in cell_par.runs:
                            run.bold = True
            doc.add_paragraph()

        elif kind == "rule":
            pass

    # ---- figures, inserted before References ----
    if figures:
        _insert_figures(doc)

    doc.save(out_path)
    return out_path


def _insert_figures(doc):
    """Append the wireframes/site map section.

    Called after the body is written; the References heading is moved back
    below the figures so the section lands where the plan specifies.
    """
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    missing = [f for f, _ in PROPOSAL_FIGURES if not (FIGURES / f).exists()]
    if missing:
        print(f"  ! figures missing, skipping image section: {missing}", file=sys.stderr)
        return

    body = doc.element.body
    # locate the References heading so we can insert before it
    ref_el = None
    for par in doc.paragraphs:
        if par.text.strip().lower() == "references" and par.style.name.startswith("Heading"):
            ref_el = par._p
            break

    anchor = []
    heading = doc.add_heading("Wireframes and Site Map", level=2)
    anchor.append(heading._p)
    for fname, caption in PROPOSAL_FIGURES:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(FIGURES / fname), width=Inches(6.5))
        anchor.append(p._p)
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(10)
        anchor.append(cap._p)

    if ref_el is not None:
        for el in anchor:
            body.remove(el)
            ref_el.addprevious(el)


# --------------------------------------------------------------------------
# HTML output (for Google Drive import)
# --------------------------------------------------------------------------

def inline_html(text):
    out = []
    for chunk, bold, italic, code, href in split_inline(text):
        esc = html_mod.escape(chunk, quote=False)
        if href:
            esc = f'<a href="{html_mod.escape(href, quote=True)}">{esc}</a>'
        if code:
            esc = f"<code>{esc}</code>"
        if italic:
            esc = f"<i>{esc}</i>"
        if bold:
            esc = f"<b>{esc}</b>"
        out.append(esc)
    return "".join(out)


def build_html(blocks, out_path, title):
    parts = [f'<html><head><meta charset="utf-8"><title>{html_mod.escape(title)}</title>'
             "</head><body>"]
    in_ul = False

    def close_ul():
        nonlocal in_ul
        if in_ul:
            parts.append("</ul>")
            in_ul = False

    for kind, payload in blocks:
        if kind != "bullet":
            close_ul()
        if kind == "heading":
            level, text = payload
            parts.append(f"<h{min(level,4)}>{inline_html(text)}</h{min(level,4)}>")
        elif kind == "para":
            parts.append(f"<p>{inline_html(payload)}</p>")
        elif kind == "refpara":
            parts.append(
                f'<p style="margin-left:36pt;text-indent:-36pt">{inline_html(payload)}</p>')
        elif kind == "bullet":
            if not in_ul:
                parts.append("<ul>")
                in_ul = True
            parts.append(f"<li>{inline_html(payload)}</li>")
        elif kind == "fence":
            parts.append(f"<pre><code>{html_mod.escape(payload)}</code></pre>")
        elif kind == "rule":
            parts.append("<hr>")
        elif kind == "table":
            rows = payload
            if not rows:
                continue
            parts.append('<table border="1" cellspacing="0" cellpadding="6" '
                         'style="border-collapse:collapse">')
            # NOTE: no <th> row. Google Drive's HTML importer renders <th> as a
            # separate body row and inserts a blank row above it. Bold <td>
            # cells import cleanly.
            for ri, row in enumerate(rows):
                cells = "".join(
                    f'<td{" bgcolor=\"#eeeeee\"" if ri == 0 else ""}>'
                    f'{"<b>" + inline_html(c) + "</b>" if ri == 0 else inline_html(c)}</td>'
                    for c in row)
                parts.append(f"<tr>{cells}</tr>")
            parts.append("</table>")
    close_ul()
    parts.append("</body></html>")
    Path(out_path).write_text("\n".join(parts))
    return out_path


# --------------------------------------------------------------------------
# Figures
# --------------------------------------------------------------------------

def find_chrome():
    for c in CHROME_CANDIDATES:
        if os.path.exists(c):
            return c
    return shutil.which("google-chrome") or shutil.which("chromium")


def regenerate_figures():
    chrome = find_chrome()
    if not chrome:
        print("Chrome/Chromium not found; cannot regenerate figures.", file=sys.stderr)
        return 1
    FIGURES.mkdir(exist_ok=True)
    for svg in sorted(DOCS.glob("*.svg")):
        try:
            root = ET.parse(svg).getroot()
            vb = root.get("viewBox", "").split()
            w, h = (int(float(vb[2])), int(float(vb[3]))) if len(vb) == 4 else (1200, 900)
        except ET.ParseError as e:
            print(f"  ! {svg.name}: {e}", file=sys.stderr)
            continue
        out = FIGURES / (svg.stem + ".png")
        subprocess.run([
            chrome, "--headless", "--disable-gpu", "--no-sandbox",
            f"--screenshot={out}", f"--window-size={w},{h}",
            "--force-device-scale-factor=2", "--default-background-color=FFFFFFFF",
            "--virtual-time-budget=3000", f"file://{svg}",
        ], capture_output=True, timeout=120)
        size = out.stat().st_size if out.exists() else 0
        print(f"  {svg.name} -> {out.name}  ({w}x{h} @2x, {size:,} bytes)")
    return 0


# --------------------------------------------------------------------------

def main():
    ap = argparse.ArgumentParser(description=__doc__,
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--figures", action="store_true", help="regenerate PNGs from the SVGs")
    ap.add_argument("--html", action="store_true", help="also emit HTML for Drive import")
    args = ap.parse_args()

    if args.figures:
        print("Regenerating figures:")
        return regenerate_figures()

    print("Building documents from Markdown:")
    for spec in DOCUMENTS:
        md_path = DOCS / spec["md"]
        if not md_path.exists():
            if not spec.get("optional"):
                print(f"  ! missing source {spec['md']}", file=sys.stderr)
            continue
        blocks = parse_blocks(md_path.read_text())
        out = build_docx(blocks, DOCS / spec["docx"], spec["title"], spec.get("figures", False))
        print(f"  {spec['md']} -> {Path(out).name}  ({Path(out).stat().st_size:,} bytes)")
        if args.html:
            hout = build_html(blocks, DOCS / spec["html"], spec["title"])
            print(f"  {spec['md']} -> {Path(hout).name}  ({Path(hout).stat().st_size:,} bytes)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
